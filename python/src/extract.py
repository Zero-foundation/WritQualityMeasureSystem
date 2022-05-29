import sys

import docx
import re
import os
from Crime import Crime
from wenshu import wenshu
from LegalBase import LegalBase
import json


def batch_extract(file_dir):
    cnt = 1
    for root, dirs, files in os.walk(file_dir):
        for name in files:
            if cnt >= 100000:
                break
            docx_path = os.path.join(root, name)
            cnt += 1
            # if cnt < 50000: continue
            if name.endswith('.docx') and not name.startswith('~$'):
                extract(docx_path, cnt // 5000)


def extract(filename, file_count):

    doc = docx.Document(filename)
    if len(doc.paragraphs) < 3:
        f = open(r'D:\PythonProcect\judgementSpider\log\parse_log', 'a+', encoding='UTF-8')
        parse_failure_log = 'failed to parse this file: ' + filename + ', file has errors.\n'
        # print(parse_failure_log)
        f.write(parse_failure_log)
        f.flush()
        f.close()
        return
    court_name = doc.paragraphs[0].text
    file_type = doc.paragraphs[1].text.replace(' ', '')
    file_number = doc.paragraphs[2].text
    content = ''

    for para in doc.paragraphs:
        content = content + '\n' + para.text
    laws = []
    ws = None
    # 刑事案件初审裁判文书
    if '刑初' in file_number:
        # print(filename)
        # example = '公诉机关xxx\n被告人xxx1\n辩护人xxx1\n辩护人xxx2\n被告人xxx2\n辩护人xxx3\n'
        # TODO: 完善全文正则
        # pattern_whole = '(?P<prosecutionOrgan>公诉机关.*)\n(?P<accused>(被告人[\s\S]*\n(辩护人.*\n)*)+)' \
        #                 '(?P<process>.*以.*起诉书.*)\n(?P<accusation>.*指控[\s\S]*\n)?(?P<facts>[\s\S]*)\n(?P<conclusion>本院认为[\s\S]*判决如下[：:])\n(?P<decision>[\s\S]*)\n(?P<rubufu>如不服.*)\n' \
        #                 '(?P<luokuan>[\s\S]*书 ?　?记 ?　?员.*)'

        pattern_law = '(((根据)|(依据)|(依照))(《.*规定))'
        pattern_crime = '(((犯)|(构成)){1}([^，。：；、]+罪){1})'
        pattern_LegalBase = '(依|根)(据|照)(《[^，。、《》]+》.*)判决如下'
        pattern_LegalBase_article = '(《[^，。、《》]+》)?(第?[零一二三四五六七八九十百千]+条)[^第、]*(第?[一二三四五六七八九]款)?(第?（[一二三四五六七八九]）项)?、?'
        anyou = ''

        pattern_whole_pre = '(?P<prosecutionOrgan>公诉机关.*)\n(?P<accused>被告((单位)|人)[\s\S]*)\n(?P<process>.*起诉书，?指控，?被告人.*)\n' \
                            '(?P<main_body>[\s\S]*)(?P<conclusion>本院认为[\s\S]*判决如下[：:])(?P<decision>[\s\S]*)\n(?P<rubufu>如不服.*)\n' \
                            '(?P<luokuan>[\s\S]*书 ?　?记 ?　?员.*)'

        res_whole = re.search(pattern_whole_pre, content)
        if (res_whole is not None):
            ws = wenshu(filename.split('\\')[-1])
            # ws.set_reply(res_whole.group('reply')) 这条放在后面单独提取
            ws.set_prosecution_organ(res_whole.group('prosecutionOrgan'))
            ws.set_accused(res_whole.group('accused'))
            ws.set_process(res_whole.group('process'))
            ws.set_facts(res_whole.group('main_body'))
            ws.set_conclusion(res_whole.group('conclusion'))
            ws.set_decision(res_whole.group('decision'))
            ws.set_rubufu(res_whole.group('rubufu'))
            ws.set_luokuan(res_whole.group('luokuan'))
            main_body = res_whole.group('main_body')
            # ws.print()

        else:
            f = open(r'D:\PythonProcect\judgementSpider\log\parse_log', 'a+', encoding='UTF-8')
            parse_failure_log = 'failed to parse this file: ' + filename + '\n'
            # print(parse_failure_log)
            f.write(parse_failure_log)
            f.flush()
            f.close()
            return
        # print(main_body)
        # TODO：提取指控内容


        # 提取适用法条
        res_legal_base = re.search(pattern_LegalBase, ws.conclusion)
        if res_legal_base is not None:
            res_article = re.findall(pattern_LegalBase_article, res_legal_base[0])
            # print(res_legal_base[0])
            # print(res_article)
            law_name = res_article[0][0]
            for article in res_article:
                if (len(article[0]) > 0):
                    law_name = article[0]
                laws.append(LegalBase(law_name, article[1], article[2], article[3]))
        for law in laws:
            ws.add_law(law.getLaw())

        # 提取案由
        res_crime = re.search(pattern_crime, ws.conclusion)
        if res_crime is not None:
            crime = res_crime[0]
            if crime[0] == '犯':
                crime = crime[1:]
            if crime[0] == '构':
                crime = crime[2:]
            ws.set_crime(crime)

        # 对未能爬取成功的被告人辩称/无异议部分进行提取
        pattern_reply = "(被告人?[^。，；\n\r]*((辩解)|(辩称)|(异议)).*\n)|(被告人?[^。，；\n]*(供述|供认).*\n)"
        res_reply = re.search(pattern_reply, content)
        if res_reply is not None:
            ws.set_reply(res_reply[0])

        # 提取被告人
        pattern_accused = '被告人(?P<name>[^。\(（,，\n]*).*\n?'
        pattern_accused_institude = '被告单位(?P<name>[^。\(（,，\n]*).*\n?'
        accused_text = ws.accused

        res_accused = re.findall(pattern_accused, accused_text)
        # print(res_accused)
        res_accused.extend(re.findall(pattern_accused_institude, accused_text))
        accused_dict = {}
        for name in res_accused:
            accused_dict[name] = {'name':'','circumstance': [], 'judge_res': [], 'crimes': []}

        # 将法院认为……部分中包含被告人姓名的部分提取，处理。先按句号分割，再按分号分割。
        conclusion_text = ws.conclusion
        # 句号分割
        conslus = conclusion_text.split('。')
        pattern_sub_accused = '(('
        for accused in res_accused:
            if '*' in accused:
                accused = str.replace(accused, '*', '\\*')
            pattern_sub_accused += '(' + accused + ')' + '|'
        pattern_sub_accused = pattern_sub_accused[0:-1] + ')、?)'
        for conclusion in conslus:
            # 分号分割
            conclusion_split_with_semicolon = conclusion.split('；')
            for conclu_semi in conclusion_split_with_semicolon:
                subed_conclu = re.sub(pattern_sub_accused, '', conclu_semi)
                # 对每个被告人姓名是否在句中做查询
                for accused in res_accused:
                    if accused in conclu_semi:
                        accused_dict[accused]['circumstance'].append(subed_conclu)
        decision_split = ws.decision.split('\n')
        p_year = '[零一二三四五六七八九十百千]*年'
        p_month = '[零一二三四五六七八九十百千]*月'
        p_day = '[零一二三四五六七八九十百千]*日'
        pattern_main_punishment = '((((管制)|(拘役)|(有期)|(有期徒刑)|(缓刑)|(剥夺政治权利))(?=([零一二三四五六七八九十百千]*年)|([零一二三四五六七八九十百千]*个?月)|([零一二三四五六七八九十百千]*[日天]))([零一二三四五六七八九十百千]*年)?([零一二三四五六七八九十百千]*个?月)?([零一二三四五六七八九十百千]*[日天])?)|(处罚金[^。，]*元)|(死刑)|(剥夺政治权利终身)|(无期徒刑)|(驱逐出境))'

        # decision中提取罪名
        for accused in res_accused:
            crime_set = set()
            for decision in decision_split:
                if '罪' in decision and accused in decision:
                    # 提取罪名
                    # for crime in crime_list:
                    #     if crime in decision:
                    #         crime_set.add(crime)
                    #         crimes[crime] += 1
                    res_crime_personal=re.findall('犯([^，。元；,.]+罪)',decision)
                    for crime_text in res_crime_personal:
                        res_crime_split = re.findall('([^、]+罪)、', crime_text)
                        if len(res_crime_split)!=0:
                            res_crime_split_afer = re.findall('、([^、]+罪)', crime_text)
                            # print(res_crime_split)
                            # print(res_crime_split_afer)
                            for element in res_crime_split_afer:
                                crime_set.add(element)
                            for element in res_crime_split:
                                crime_set.add(element)
                        else:
                            crime_set.add(crime_text)


            for key in crime_set:
                accused_dict[accused]['crimes'].append(key)
        # 在decision中提取结果
        for accused in res_accused:
            for decision in decision_split:
                if accused in decision:
                    # 提取结果
                    res_final_punishment = re.search('(决定执行.*)', decision)
                    if res_final_punishment is not None:
                        res_punishment = re.findall(pattern_main_punishment, res_final_punishment[0])
                        if len(res_punishment) > 0:
                            for punishment in res_punishment:
                                accused_dict[accused]['judge_res'].append(punishment[0])
                    else:
                    # 无刑罚合并情况（未执行完/数罪并罚）
                        res_punishment = re.findall(pattern_main_punishment, decision)
                        if len(res_punishment) > 0:
                            for punishment in res_punishment:
                                accused_dict[accused]['judge_res'].append(punishment[0])
                    # 数罪并罚的应当需要特殊处理
        keys=list(accused_dict.keys())
        accused_array=[]
        for accused in keys:
            accused_dict[accused]["name"]=accused
            if len(accused_dict[accused]['judge_res'])==0:
                accused_dict.pop(accused)
        keys = list(accused_dict.keys())
        for accused in keys:
            accused_array.append(accused_dict[accused])

        #将accused_dict变为Object的array

        ws.set_accused_list(accused_array)
        # print(accused_dict)
        # file_path = 'D:\\PythonProcect\\judgementSpider\\data\\alter\\' + 'wenshu' + str(file_count) + '.json'
        # f = open(file_path, 'a+', encoding='UTF-8')
        # json.dump(ws.getJson(), open(file_path, 'a+', encoding='UTF-8'), indent=4, ensure_ascii=False)
        # f.write(',')
        return ws.getJson()



if __name__ == '__main__':
    fileName=sys.argv[1]
    res=extract(fileName,0)
    print(res)
    # filename1 = r"D:\PythonProcect\judgementSpider\judgements\董某某故意伤害罪一审刑事判决书.docx"
    # filename2 = r"D:\PythonProcect\judgementSpider\judgements\王林危险驾驶罪刑事一审刑事判决书.docx"
    # extract(filename1)
    # extract(filename2)
    # extract(r'D:\PythonProcect\xingchu\68南京胜科水务有限公司ＺＨＥＮＧＱＩＡＯＧＥＮＧ等污染环境罪一审刑事判决书.docx')
    # extract(r'D:\PythonProcect\xingchu\1037佟毫故意伤害罪一审刑事判决书.docx', 0)
    # extract(r'D:\PythonProcect\xingchu\106蒋珺盗窃罪一审刑事判决书.docx',1)
    # batch_extract(r'D:\PythonProcect\xingchu')
    # save_crime()
