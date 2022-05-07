import os

import shutil
import docx
from win32com import client as wc


def parse(file_dir):
    flag_pass = True
    for root, dirs, files in os.walk(file_dir):
        target_dir = r'D:\PythonProcect\xingchu'
        for name in files:

            docx_path = os.path.join(root, name)

            if (flag_pass):
                continue

            if name.endswith('.doc') and '判决' in name and not name.startswith('~$'):
                print('processing ' + docx_path)
                # shutil.copyfile(docx_path,os.path.join(target_dir,name))
                word = wc.Dispatch("Word.Application")
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(docx_path + 'x', 12)
                doc.Close()

                # processing D:\PythonProcect\judgements\丁宝宏猥亵儿童罪一审刑事判决书.doc
                print('copying')
                shutil.copyfile(docx_path + 'x', os.path.join(target_dir, name) + 'x')
                # doc = docx.Document(docx_path+'x')
                # if(len(doc.paragraphs)>=3):
                #     file_number = doc.paragraphs[2].text
                #     if '刑初' in file_number:
                #         shutil.copyfile(docx_path + 'x', os.path.join(target_dir, name) + 'x')

            # print("段落数:" + str(len(file.paragraphs)))  # 输出段落数


def remove_dumplicate():
    file_dir = r'D:\PythonProcect\xingchu'
    dic = {}
    list = []
    for root, dirs, files in os.walk(file_dir):
        hit_cnt = 0
        not_hit_cnt = 0
        for name in files:
            docx_path = os.path.join(root, name)
            doc = docx.Document(docx_path)
            if (len(doc.paragraphs) >= 3):
                file_number = doc.paragraphs[2].text
                if file_number not in dic.keys():
                    dic[file_number] = docx_path
                    hit_cnt += 1
                    print("hit: " + str(hit_cnt))
                else:
                    list.append(docx_path)
                    print(docx_path)
                    not_hit_cnt += 1
                    print("not hit: " + str(not_hit_cnt))
    print(list)
    for path in list:
        os.remove(path)


if __name__ == '__main__':
    # parse(r"D:\PythonProcect\judgements")
    remove_dumplicate()
